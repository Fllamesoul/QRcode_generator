import imp
from tkinter import *
import qrcode
#import PIL
from docx import Document
import os
from docx.shared import Cm, Pt
import imp


root = Tk()
root.title('Генератор qr-кода')
root.geometry('300x350')


# button
butt_generate = Button(root, text = 'Сделать все')
butt_generate.place(rely = 0.3, relx = 0.5, anchor = N)
butt_generate.bind('<Button-1>', lambda event: start())
###

# count
COUNT = StringVar()
area_COUNT = Entry(root, font = ("Helvetica", 12), justify = 'center', bd=0, state='readonly', textvariable = COUNT)
area_COUNT.place(rely = 0.5, relx = 0.5, anchor = N)
area_COUNT.bind("<Button-3>")
###


fp, pathname, description = imp.find_module("excel")
_mod = imp.load_module("excel", fp, pathname, description)

PIC_FILE_NAME = "pic.png"
NUMBER_COLS = 5 * 2
NUMBER_ROWS = 0
TABLE_CELL_WIDTH = Cm(5)
WIDTH_QR = Cm(1)
HEIGHT_QR = Cm(1)
FONT_SIZE = Pt(5)

path_excel_file = "Test table.xlsx"
list_num_col_for_QR =   (1, 2, 3, 4, 6, 7, 8, 9)
list_num_col_QR_descr = (4, 7, 8, 9)
path_doc = "QR_codes.docx"
path_template_docx_file = "Template.docx"         
path_excel_file = "Для инвентаризации.xlsx"       
rezult_doc_path = "result.docx"


# функция создает таблицу в формате ".doc" с QR-кодами
# и описанием мат. технических ресурсов
def make_QR_doc_list(excel, table, list_col_for_QR, list_col_QR_descr):
    num_row_doc = 0
    num_col_doc = 0
    for count in range(10): #FIXME 25!!!
        #добавление строк
        if num_col_doc >= NUMBER_COLS:
            num_col_doc = 0
            num_row_doc += 1

        # добавление содержимого
        if not num_col_doc % 2:
            make_img_QR(excel, count + 2, list_col_for_QR) 
            run = table.cell(num_row_doc, num_col_doc).paragraphs[0].add_run()
            run.add_picture(PIC_FILE_NAME, width = WIDTH_QR, height = HEIGHT_QR)
        else:
            paragraph = table.cell(num_row_doc, num_col_doc).paragraphs[0]
            paragraph.style.font.size = FONT_SIZE
            run = paragraph.add_run()
            run.add_text(get_string_from_excel(excel, count, list_col_QR_descr))
        num_col_doc += 1
        
    os.remove(PIC_FILE_NAME)


def get_string_from_excel(excel_file, num_row, list_num_col):
    num_row += 2
    data = ""
    for i in list_num_col:
        data += excel_file.get_data_cell(num_row, i)
    return data


def make_img_QR(excel_file, num_row, list_num_col_for_QR):
    str_QR = get_string_from_excel(excel_file, num_row, list_num_col_for_QR)
    img = qrcode.make(str_QR)
    img.save(PIC_FILE_NAME)


def set_rows_height(table, height):
    for row in table.rows:
        row.height = height

#################################################################################


def start():
    fp, pathname, description = imp.find_module("excel")
    _mod = imp.load_module("excel", fp, pathname, description)
    ex = _mod.Excel(path_excel_file)
    doc_QR_document = Document(path_template_docx_file)
    table = doc_QR_document.tables[0]
    make_QR_doc_list(ex, table, list_num_col_for_QR, list_num_col_QR_descr)
    doc_QR_document.save(rezult_doc_path)
    COUNT.set("файл rezult.docx cохранен")


root.mainloop()



def get_velues_from_excel(excel_file, num_row, list_num_col):
    data = []
    for i in list_num_col:
        data.append(excel_file.get_data_cell(num_row, i))
    return data

